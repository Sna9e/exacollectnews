import streamlit as st
import sys
import json
import os
import datetime
import concurrent.futures
import platform
import difflib
from typing import List

# ================= 0. 核心库引用 =================
from pydantic import BaseModel, Field
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI  
from exa_py import Exa # 🔴 新增：Exa AI 官方 SDK

# ================= 1. 核心网络配置 =================
if platform.system() == "Windows":
    os.environ["http_proxy"] = "http://127.0.0.1:7890"
    os.environ["https_proxy"] = "http://127.0.0.1:7890"
else:
    os.environ.pop("http_proxy", None)
    os.environ.pop("https_proxy", None)

# ================= 2. 文档库引用 =================
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="DeepSeek 科技探员", page_icon="🐳", layout="wide")

# ================= 3. 定义结构化数据 =================
class NewsItem(BaseModel):
    title: str = Field(description="新闻标题（务必翻译为中文）")
    source: str = Field(description="来源媒体（保留原名）")
    date_check: str = Field(description="严格核实新闻发生的真实日期，格式 YYYY-MM-DD。")
    summary: str = Field(description="约300字的深度商业分析。必须严格分段并带有标识：【事件核心】、【深度细节/数据支撑】、【行业深远影响】。")
    importance: int = Field(description="重要性 1-5")

class NewsReport(BaseModel):
    news: List[NewsItem] = Field(description="新闻列表")

# ================= 4. 内置 DeepSeek 驱动 =================
class EnterpriseDeepSeekDriver:
    def __init__(self, api_key, model_id):
        self.valid = False
        if not api_key: return
        try:
            self.client = OpenAI(api_key=api_key, base_url="https://api.deepseek.com")
            self.model_id = model_id
            self.valid = True
        except Exception:
            pass

    def analyze_structural(self, prompt, structure_class):
        if not self.valid: return None
        schema_str = json.dumps(structure_class.model_json_schema(), ensure_ascii=False)
        sys_prompt = f"你是顶级商业情报分析师。必须严格按此 JSON Schema 返回数据，不带任何废话：\n{schema_str}"
        try:
            response = self.client.chat.completions.create(
                model=self.model_id,
                messages=[{"role": "system", "content": sys_prompt}, {"role": "user", "content": prompt}],
                response_format={"type": "json_object"},
                temperature=0.1, 
                max_tokens=4096 
            )
            raw_text = response.choices[0].message.content.strip()
            try:
                json_obj = json.loads(raw_text)
                if isinstance(json_obj, list): json_obj = {"news": json_obj}
                return structure_class(**json_obj)
            except Exception: return None
        except Exception: return None

# ================= 5. 核心业务函数 =================

# 🔴 史诗级进化：使用 Exa AI 一次性完成“神级搜索 + 正文提取”
def search_and_extract_with_exa(query, sites_text, time_opt, exa_key, max_results=10):
    if not exa_key: return "", 0, []
    
    exa = Exa(api_key=exa_key)
    sites = [s.strip() for s in sites_text.split('\n') if s.strip()]
    
    # 将时间选项转化为 Exa 支持的 published_date 过滤
    start_date = None
    if time_opt == "d":
        start_date = (datetime.datetime.now() - datetime.timedelta(days=2)).strftime("%Y-%m-%d")
    elif time_opt == "w":
        start_date = (datetime.datetime.now() - datetime.timedelta(days=7)).strftime("%Y-%m-%d")
    elif time_opt == "m":
        start_date = (datetime.datetime.now() - datetime.timedelta(days=30)).strftime("%Y-%m-%d")

    # 构建强大的 Exa 查询参数
    search_args = {
        "query": f"latest news about {query}", 
        "type": "auto", 
        "use_autoprompt": True, # 开启 Exa 自动扩写 Prompt 增强语义
        "num_results": max_results,
        "contents": {"text": {"max_characters": 6000}} # 🔴 直接提取前6000字正文，抛弃爬虫！
    }
    
    if sites:
        search_args["include_domains"] = sites
    if start_date:
        search_args["start_published_date"] = start_date

    try:
        response = exa.search(**search_args)
        
        full_content = ""
        valid_count = 0
        links = []
        
        for result in response.results:
            if result.text:
                valid_count += 1
                links.append({'href': result.url})
                # 直接将 Exa 提取的干净正文喂给大模型
                full_content += f"\n\n=== SOURCE START: {result.url} ===\n{result.text}\n=== SOURCE END ===\n"
                
        return full_content, valid_count, links
    except Exception as e:
        print(f"Exa API Error: {e}")
        return "", 0, []

def map_reduce_analysis(ai_driver, topic, full_text, current_date, time_opt):
    if not full_text or len(full_text) < 100: return []
    docs = RecursiveCharacterTextSplitter(chunk_size=8000, chunk_overlap=1000).create_documents([full_text])
    all_extracted_news = []

    def process_single_doc(doc):
        map_prompt = f"""
        【全局时间锚点】：今天是 **{current_date}**。
        要求的时间范围是：【{time_opt}】。
        任务：从以下文本提取关于【{topic}】的新闻情报。
        红线：
        1. 严格时间审查：发现发生时间早于【{time_opt}】之前（如几月前、去年），直接丢弃！
        2. 【{topic}】必须是绝对主角！
        无符合条件的内容必须返回 `{{"news": []}}`。
        文本：{doc.page_content}
        """
        return ai_driver.analyze_structural(map_prompt, NewsReport)

    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        for future in concurrent.futures.as_completed([executor.submit(process_single_doc, d) for d in docs]):
            res = future.result()
            if res and res.news: all_extracted_news.extend(res.news)

    if not all_extracted_news: return []
    combined_json = json.dumps([item.model_dump() for item in all_extracted_news], ensure_ascii=False)

    reduce_prompt = f"""
        【全局时间锚点】：今天是 **{current_date}**。
        你是极其严苛的科技媒体总编。
        任务：
        1. 终极时间清洗：任何陈年旧闻，全部无情删掉！
        2. 合并去重：报道同一事件的新闻必须合并。
        3. 深度扩写与高级排版：将每条新闻的 summary 扩展至 300 字左右。必须在 summary 中使用明显的分段和换行，明确包含以下三个部分：
           【事件核心】：概括事件
           【深度细节】：核心数据与细节支撑
           【行业影响】：精简的行业深远影响
        4. 按重要性降序，最多保留最核心的 5 条。
        数据：{combined_json}
    """
    final_report = ai_driver.analyze_structural(reduce_prompt, NewsReport)
    return final_report.news if final_report else []

def generate_word(data, filename, model_name):
    doc = Document()
    normal_style = doc.styles['Normal']
    normal_style.font.name = '微软雅黑'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    normal_style.font.size = Pt(10.5) 
    for i in range(1, 4):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = '微软雅黑'
        h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if i == 1:
            h_style.font.color.rgb = RGBColor(0, 51, 102)

    title = doc.add_heading("DeepSeek 企业级深度科技研报", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(f"生成日期: {datetime.date.today()}  |  数据来源: Exa AI 神经搜索引擎  |  分析模型: {model_name}")
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    meta_run.font.size = Pt(9)
    
    doc.add_paragraph("━" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for section in data:
        doc.add_heading(f"🔷 专题：{section['topic']}", level=1)
        if not section['data']:
            doc.add_paragraph("    在指定时间范围内，未发现符合标准的重大情报。").font.italic = True
            continue
            
        for news in section['data']:
            doc.add_heading(f"🔹 {news.title}", level=2)
            p_info = doc.add_paragraph()
            run_info = p_info.add_run(f"    📌 来源: {news.source}    |    🕒 时间: {news.date_check}    |    🔥 价值评级: {'⭐'*news.importance}")
            run_info.font.color.rgb = RGBColor(100, 100, 100)
            run_info.font.bold = True
            
            p_summary = doc.add_paragraph(news.summary)
            p_summary.paragraph_format.line_spacing = 1.5 
            p_summary.paragraph_format.first_line_indent = Pt(21) 
            
            divider = doc.add_paragraph("┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈┈")
            divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
            divider.runs[0].font.color.rgb = RGBColor(200, 200, 200)
    
    path = f"{filename}.docx"
    doc.save(path)
    return path

# ================= 6. 主界面 =================
with st.sidebar:
    st.header("🐳 DeepSeek 控制台")
    api_key = st.text_input("DeepSeek API Key", type="password")
    
    # 🔴 更换为 Exa API Key
    exa_key = st.text_input("Exa AI API Key (必填)", type="password", help="去 dashboard.exa.ai 免费获取此 Key，体验神级语义搜索与正文提取！")
    
    model_id = st.selectbox("模型", ["deepseek-chat"], index=0)
    st.divider()
    
    time_opt = st.selectbox("时间范围（绝对严控）", ["过去 24 小时", "过去 1 周", "过去 1 个月", "不限时间"], index=0)
    time_limit_dict = {"过去 24 小时": "d", "过去 1 周": "w", "过去 1 个月": "m", "不限时间": None}
    
    st.markdown("**垂直情报源雷达**")
    sites = st.text_area("重点搜索源", "techcrunch.com\ntheverge.com\nengadget.com\ncnet.com\nbloomberg.com/technology\nelectrek.co\ninsideevs.com\nroadtovr.com\nuploadvr.com\n36kr.com\nithome.com\nhuxiu.com\ngeekpark.net\nvrtuoluo.cn\nd1ev.com", height=250)
    file_name = st.text_input("文件名", f"深度研报_{datetime.date.today()}")

st.title("🐳 企业情报探员 (Exa 智能体版)")
query_input = st.text_input("输入主题 (用 \\ 隔开，外媒源建议用英文如：Google \\ Apple)", "Tesla Robotaxi \\ Apple Vision Pro")
btn = st.button("🚀 开始生成研报", type="primary")

if btn:
    if not api_key or not exa_key:
        st.error("❌ 请先在左侧边栏填入 DeepSeek 和 Exa 的 API Key！")
    elif not query_input:
        st.warning("请输入关键词！")
    else:
        topics = [t.strip() for t in query_input.split('\\') if t.strip()]
        all_data = []
        ai = EnterpriseDeepSeekDriver(api_key, model_id)
        current_date_str = datetime.date.today().strftime("%Y年%m月%d日")
        
        global_seen_titles = []

        st.info("🚀 探员已出击，Exa 神经元网络正在提取全球情报...")

        for topic in topics:
            st.markdown(f"#### 🔵 追踪目标: 【{topic}】 (要求: {time_opt})")
            
            with st.spinner(f"正在全网智能嗅探并直抽正文... (无需爬虫)"):
                # 🔴 直接一步到位，获取满血正文！
                full_text_data, valid_count, links = search_and_extract_with_exa(topic, sites, time_limit_dict[time_opt], exa_key)
            
            if not full_text_data: 
                st.warning(f"⚠️ {topic}：未搜寻到任何有效新闻。说明目标近期很安静或 Exa Key 有误！")
                continue
                
            st.write(f"🔍 成功获取并提取了 {valid_count} 个高价值网页的核心正文。DeepSeek 正在执行深度分析...")

            with st.spinner("AI 正在剔除旧闻与重复项，撰写商业分析..."):
                final_news_list = map_reduce_analysis(ai, topic, full_text_data, current_date_str, time_opt)
            
            if final_news_list:
                deduped_news = []
                for news in final_news_list:
                    is_duplicate = False
                    for seen_title in global_seen_titles:
                        similarity = difflib.SequenceMatcher(None, news.title, seen_title).ratio()
                        if similarity > 0.6:
                            is_duplicate = True
                            break
                    
                    if not is_duplicate:
                        deduped_news.append(news)
                        global_seen_titles.append(news.title)
                
                if deduped_news:
                    all_data.append({"topic": topic, "data": deduped_news})
                    filtered_count = len(final_news_list) - len(deduped_news)
                    st.success(f"✅ 【{topic}】分析完毕！已锁定 {len(deduped_news)} 条新鲜情报。" + 
                               (f"(已跨主题去重过滤 {filtered_count} 条重复事件)" if filtered_count > 0 else ""))
                else:
                    st.warning(f"⚠️ 【{topic}】提炼出的新闻均与之前的主题高度重合，已执行全局去重抹杀！")
                    
            else:
                st.warning(f"⚠️ 【{topic}】搜到的网页经 AI 严格审判，全被判定为旧闻或非核心新闻，已执行抹杀过滤。")
            
            st.divider()

        if all_data:
            path = generate_word(all_data, file_name, model_id)
            st.balloons()
            st.success("🎉 全链条任务执行完毕！请下载查看由 Exa+DeepSeek 联合驱动的高级排版研报。")
            with open(path, "rb") as f:
                st.download_button("📥 立即下载精美排版研报 (Word)", f, file_name=path, type="primary")
        else:
            st.error(f"❌ 任务结束。在严格的时效与去重约束下，所有关键词均未产生独立且有效的大事件情报。")
